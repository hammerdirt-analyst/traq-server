from __future__ import annotations

"""Report-letter text/prompt generation and PDF rendering utilities.

Authors:
    Roger Erismann (https://hammerdirt.solutions), OpenAI Codex

Purpose:
    Build a client-facing narrative letter from reviewed form/transcript data
    and render it to PDF for final delivery.

Document/layout contract:
    - Paper size: US Letter (8.5 x 11 in)
    - Margins: 1.0 in on all sides
    - Body font: Times-Roman, target 11 pt (auto-reduced when needed to keep
      letter/signature on one page)
    - Heading font: Times-Bold 12 pt for address block first lines
    - Subject prefix: "Subject :" in bold
"""

from dataclasses import dataclass
from datetime import datetime
from typing import Iterable, Sequence

import json
import os

from openai import OpenAI


@dataclass(frozen=True)
class ReportLetterInputs:
    """Normalized inputs used by plain-text letter rendering."""

    user_name: str
    correspondence: Sequence[str]
    billing_lines: Sequence[str]
    job_location: str
    subject_line: str
    date_text: str
    paragraphs: Sequence[str]
    user_isa_number: str
    user_phone: str
    user_email: str


def build_report_letter(
    *,
    profile: dict | None,
    job: dict | None,
    summary: str,
    form_data: dict | None = None,
    date_text: str | None = None,
) -> str:
    """Create normalized report-letter text from profile/job/summary data.

    Args:
        profile: User profile payload.
        job: Job metadata payload.
        summary: Five-paragraph summary text.
        form_data: Form payload, used to derive subject fields.
        date_text: Optional date override.

    Returns:
        Plain-text letter body used by PDF renderer.
    """
    profile = profile or {}
    job = job or {}
    form_data = form_data or {}
    inputs = ReportLetterInputs(
        user_name=_as_text(profile.get("name")),
        correspondence=_build_user_address_block(profile),
        billing_lines=_build_customer_address_block(job),
        job_location=_as_text(
            job.get("job_address")
            or job.get("address")
            or job.get("address_tree_location")
        ),
        subject_line=_build_subject_line(job, form_data),
        date_text=date_text or _format_date(),
        paragraphs=_extract_paragraphs(summary),
        user_isa_number=_as_text(profile.get("isa_number")),
        user_phone=_as_text(profile.get("phone")),
        user_email=_as_text(profile.get("correspondence_email")),
    )
    inputs = _apply_format_rules(inputs)
    return render_report_letter(inputs)


def render_report_letter(inputs: ReportLetterInputs) -> str:
    """Render normalized letter inputs to a single plain-text letter."""
    paragraphs = list(inputs.paragraphs)
    while len(paragraphs) < 5:
        paragraphs.append("")
    paragraph_five = paragraphs[4]
    billing_block = _indent_block(inputs.billing_lines, indent=40)

    lines: list[str] = []
    lines.append(inputs.user_name or "")
    lines.extend(inputs.correspondence or [""])
    lines.append("")
    if billing_block:
        lines.append("")
        lines.extend(billing_block)
        lines.append("")
    lines.append(inputs.date_text)
    lines.append("")
    subject = inputs.subject_line or inputs.job_location
    lines.append(f"Subject : {subject}".strip())
    lines.append("")
    lines.append(paragraphs[0])
    lines.append("")
    lines.append(paragraphs[1])
    lines.append("")
    lines.append(paragraphs[2])
    lines.append("")
    lines.append(paragraphs[3])
    lines.append("")
    lines.append(paragraph_five)
    lines.append("")
    lines.append(inputs.user_name or "")
    if inputs.user_isa_number:
        lines.append(f"ISA - {inputs.user_isa_number}")
    if inputs.user_phone:
        lines.append(inputs.user_phone)
    if inputs.user_email:
        lines.append(inputs.user_email)
    return "\n".join(lines).strip() + "\n"


def generate_report_letter_pdf(
    letter_text: str,
    output_path: str,
    *,
    sender_name: str | None = None,
    customer_name: str | None = None,
    signature_name: str | None = None,
    signature_isa: str | None = None,
    job_number: str | None = None,
    report_images: Sequence[dict[str, str]] | None = None,
) -> None:
    """Render report letter text to PDF, including optional job images.

    Layout settings:
        - Page size: US Letter (`reportlab.lib.pagesizes.letter`)
        - Margins: left/right/top/bottom = 1.0 inch
        - Main body font: Times-Roman, base 11 pt (auto-fit to 11/10/9/8)
        - Name headings: Times-Bold 12 pt for top sender/recipient lines
        - Footer: 8 pt note referencing the TRAQ form/job number

    Header behavior:
        - Sender block is placed top-left.
        - Recipient block is placed top-right at the same vertical origin.

    Pagination behavior:
        - Letter body is kept on one page when possible by modest font-size
          reduction.
        - Images, when present, are rendered after the letter section and may
          extend to additional pages.
    """
    try:
        from reportlab.lib.pagesizes import letter  # type: ignore
        from reportlab.lib.units import inch  # type: ignore
        from reportlab.lib.utils import ImageReader  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "Report letter PDF generation requires reportlab. Install it in the server environment."
        ) from exc

    page_width, page_height = letter
    left_margin = inch
    right_margin = inch
    top_margin = inch
    bottom_margin = inch
    max_width = page_width - left_margin - right_margin

    pdf = canvas.Canvas(output_path, pagesize=letter)
    font_name = "Times-Roman"
    bold_font_name = "Times-Bold"
    base_font_size = 11
    heading_font_size = 12
    y = page_height - top_margin

    lines = letter_text.splitlines()
    normalized_sender = (sender_name or "").strip()
    normalized_customer = (customer_name or "").strip()
    normalized_signature_name = (signature_name or "").strip()
    normalized_signature_isa = (signature_isa or "").strip()

    sender_indices = [
        idx for idx, line in enumerate(lines) if line.strip() == normalized_sender
    ] if normalized_sender else []
    signature_name_indices = [
        idx for idx, line in enumerate(lines) if line.strip() == normalized_signature_name
    ] if normalized_signature_name else sender_indices
    customer_indices = [
        idx for idx, line in enumerate(lines) if line.strip() == normalized_customer
    ] if normalized_customer else []

    first_address_idx = sender_indices[0] if sender_indices else None
    billing_first_idx = customer_indices[0] if customer_indices else None
    subject_idx = next(
        (idx for idx, line in enumerate(lines) if line.strip().startswith("Subject :")),
        None,
    )
    signature_name_idx = signature_name_indices[-1] if signature_name_indices else None
    signature_isa_idx = next(
        (idx for idx, line in enumerate(lines) if line.strip() == normalized_signature_isa),
        None,
    ) if normalized_signature_isa else next(
        (idx for idx, line in enumerate(lines) if line.strip().startswith("ISA -")),
        None,
    )

    def _new_page() -> None:
        nonlocal y
        pdf.showPage()
        y = page_height - top_margin

    def _collect_block(start_idx: int | None) -> tuple[list[str], set[int]]:
        if start_idx is None:
            return [], set()
        block: list[str] = []
        used: set[int] = set()
        for i in range(start_idx, len(lines)):
            if not lines[i].strip():
                used.add(i)
                break
            block.append(lines[i].strip())
            used.add(i)
        return block, used

    sender_block, sender_used = _collect_block(first_address_idx)
    customer_block, customer_used = _collect_block(billing_first_idx)
    header_used = sender_used | customer_used

    # Sender (left) and recipient (right) at the same top level.
    if sender_block or customer_block:
        left_y = y
        right_y = y
        for i, line in enumerate(sender_block):
            lf = bold_font_name if i == 0 else font_name
            ls = heading_font_size if i == 0 else base_font_size
            pdf.setFont(lf, ls)
            pdf.drawString(left_margin, left_y, line)
            left_y -= ls * 1.2
        for i, line in enumerate(customer_block):
            rf = bold_font_name if i == 0 else font_name
            rs = heading_font_size if i == 0 else base_font_size
            pdf.setFont(rf, rs)
            text_w = pdf.stringWidth(line, rf, rs)
            pdf.drawString(page_width - right_margin - text_w, right_y, line)
            right_y -= rs * 1.2
        y = min(left_y, right_y) - base_font_size

    footer_note = f"See TRAQ Form job {job_number}" if job_number else "See TRAQ Form"
    footer_size = 8
    footer_reserve = footer_size * 2.2

    def _line_height(idx: int, raw_line: str, fs: int) -> float:
        if not raw_line.strip():
            return fs * 1.4
        if idx == subject_idx:
            prefix = "Subject :"
            stripped = raw_line.strip()
            suffix = stripped[len(prefix):].strip() if stripped.startswith(prefix) else stripped
            pdf.setFont(bold_font_name, fs)
            prefix_width = pdf.stringWidth(f"{prefix} ", bold_font_name, fs)
            available_first = max_width - prefix_width
            wrapped = _wrap_line(
                pdf,
                suffix,
                available_first if available_first > 0 else max_width,
                font_name,
                fs,
            ) if suffix else [""]
            return fs * 1.2 * max(1, len(wrapped))
        wrapped = _wrap_line(pdf, raw_line, max_width, font_name, fs)
        return fs * 1.2 * max(1, len(wrapped))

    body_indices = [i for i in range(len(lines)) if i not in header_used]
    font_size = base_font_size
    for fs in (11, 10, 9, 8):
        needed = sum(_line_height(i, lines[i], fs) for i in body_indices)
        if needed <= (y - bottom_margin - footer_reserve):
            font_size = fs
            break

    for idx, raw_line in enumerate(lines):
        if idx in header_used:
            continue
        if not raw_line.strip():
            y -= font_size * 1.4
            if y <= bottom_margin + footer_reserve:
                _new_page()
            continue

        line_font = font_name
        line_size = font_size
        if idx == signature_name_idx or idx == signature_isa_idx:
            line_font = bold_font_name

        if idx == subject_idx:
            stripped = raw_line.strip()
            prefix = "Subject :"
            suffix = stripped[len(prefix):].strip() if stripped.startswith(prefix) else stripped
            pdf.setFont(bold_font_name, font_size)
            prefix_width = pdf.stringWidth(f"{prefix} ", bold_font_name, font_size)
            available_first = max_width - prefix_width
            value_lines = _wrap_line(
                pdf,
                suffix,
                available_first if available_first > 0 else max_width,
                font_name,
                font_size,
            ) if suffix else [""]
            first_value = value_lines[0] if value_lines else ""
            pdf.drawString(left_margin, y, prefix)
            pdf.setFont(font_name, font_size)
            if first_value:
                pdf.drawString(left_margin + prefix_width, y, first_value)
            y -= font_size * 1.2
            if y <= bottom_margin + footer_reserve:
                _new_page()
            for continuation in value_lines[1:]:
                pdf.setFont(font_name, font_size)
                pdf.drawString(left_margin + prefix_width, y, continuation)
                y -= font_size * 1.2
                if y <= bottom_margin + footer_reserve:
                    _new_page()
            continue

        for line in _wrap_line(pdf, raw_line, max_width, line_font, line_size):
            pdf.setFont(line_font, line_size)
            pdf.drawString(left_margin, y, line)
            y -= line_size * 1.2
            if y <= bottom_margin + footer_reserve:
                _new_page()

    # Footer note.
    pdf.setFont(font_name, footer_size)
    pdf.drawString(left_margin, bottom_margin - (footer_size * 0.2), footer_note)

    images = list(report_images or [])
    if images:
        section_gap = font_size * 1.6
        heading_gap = font_size * 1.4
        if y - (section_gap + heading_gap) <= bottom_margin:
            _new_page()
        else:
            y -= section_gap
        pdf.setFont(bold_font_name, heading_font_size)
        pdf.drawString(left_margin, y, "Images")
        y -= heading_gap
        for idx, item in enumerate(images[:5], start=1):
            image_path = (item.get("path") or "").strip()
            if not image_path:
                continue
            try:
                reader = ImageReader(image_path)
                img_w, img_h = reader.getSize()
            except Exception:
                continue
            if img_w <= 0 or img_h <= 0:
                continue
            caption = (item.get("caption") or "").strip() or f"Photo {idx}"
            caption_text = f"{idx}. {caption}"
            caption_lines = _wrap_line(pdf, caption_text, max_width, font_name, font_size)
            max_img_h = (page_height - top_margin - bottom_margin) * 0.42
            scale = min(max_width / float(img_w), max_img_h / float(img_h), 1.0)
            draw_w = float(img_w) * scale
            draw_h = float(img_h) * scale
            required = draw_h + (len(caption_lines) * font_size * 1.2) + font_size
            if y - required <= bottom_margin:
                _new_page()
            pdf.drawImage(
                reader,
                left_margin,
                y - draw_h,
                width=draw_w,
                height=draw_h,
                preserveAspectRatio=True,
                mask="auto",
            )
            y -= draw_h + (font_size * 0.8)
            pdf.setFont(font_name, font_size)
            for line in caption_lines:
                pdf.drawString(left_margin, y, line)
                y -= font_size * 1.2
            y -= font_size * 0.6
    pdf.showPage()
    pdf.save()


def generate_report_letter_docx(
    letter_text: str,
    output_path: str,
) -> None:
    """Render the generated letter to DOCX for word-processor editing."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "Report letter DOCX generation requires python-docx. Install it in the server environment."
        ) from exc

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    for raw_line in letter_text.splitlines():
        document.add_paragraph(raw_line)

    document.save(output_path)


def _apply_format_rules(inputs: ReportLetterInputs) -> ReportLetterInputs:
    """Apply letter-format normalization rules to prepared inputs."""
    correspondence = _strip_state_zip_from_lines(inputs.correspondence)
    billing_lines = _strip_state_zip_from_lines(inputs.billing_lines)
    job_location = _strip_state_zip_from_subject(inputs.job_location)
    return ReportLetterInputs(
        user_name=inputs.user_name,
        correspondence=correspondence,
        billing_lines=billing_lines,
        job_location=job_location,
        subject_line=_strip_state_zip_from_subject(inputs.subject_line),
        date_text=inputs.date_text,
        paragraphs=inputs.paragraphs,
        user_isa_number=inputs.user_isa_number,
        user_phone=inputs.user_phone,
        user_email=inputs.user_email,
    )


def _as_text(value: object | None) -> str:
    """Convert optional values to stripped string content."""
    if value is None:
        return ""
    text = str(value).strip()
    return text


def _compact_lines(lines: Iterable[str]) -> list[str]:
    """Drop blank/whitespace-only lines."""
    return [line for line in lines if line.strip()]


def _indent_block(lines: Sequence[str], indent: int = 40) -> list[str]:
    """Left-pad lines by fixed spaces (legacy helper)."""
    if not lines:
        return []
    prefix = " " * max(indent, 0)
    return [f"{prefix}{line}" if line else "" for line in lines]


def _extract_paragraphs(summary: str) -> list[str]:
    """Split summary text into paragraph chunks."""
    summary = summary.strip()
    if not summary:
        return []
    raw = [chunk.strip() for chunk in summary.split("\n\n")]
    paragraphs = [chunk for chunk in raw if chunk]
    if len(paragraphs) <= 1:
        paragraphs = [line.strip() for line in summary.split("\n") if line.strip()]
    return paragraphs


def _wrap_line(
    pdf_canvas: "canvas.Canvas",
    text: str,
    max_width: float,
    font_name: str,
    font_size: int,
) -> list[str]:
    """Word-wrap text to fit width for the requested font on the canvas."""
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if pdf_canvas.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _truncate_line_to_width(
    pdf_canvas: "canvas.Canvas",
    text: str,
    max_width: float,
    font_name: str,
    font_size: int,
) -> str:
    """Truncate a line with ellipsis to fit max width (diagnostic helper)."""
    if pdf_canvas.stringWidth(text, font_name, font_size) <= max_width:
        return text
    truncated = text
    while truncated and pdf_canvas.stringWidth(
        f"{truncated}…",
        font_name,
        font_size,
    ) > max_width:
        truncated = truncated[:-1]
    return f"{truncated}…"


def polish_summary(
    summary: str,
    *,
    form_data: dict | None = None,
    transcript: str | None = None,
) -> str:
    """Polish generated summary text with LLM without adding new facts."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return summary
    model = os.environ.get("TRAQ_OPENAI_SUMMARY_MODEL", "gpt-4o-mini")
    client = OpenAI(api_key=api_key)

    system_prompt = (
        "You are a professional arborist report editor. "
        "Polish the provided summary so it reads cleanly and professionally. "
        "Do not add new facts. Preserve all stated facts and measurements. "
        "Keep it in paragraph form without bullets."
        "Do not add any conclusions or recommendations."
    )
    context = ""
    if form_data:
        context += f"Extracted data (JSON):\\n{json.dumps(form_data, ensure_ascii=True)}\\n\\n"
    if transcript:
        context += f"Transcript:\\n{transcript}\\n\\n"
    user_prompt = (
        f"{context}Summary to polish:\\n{summary}\\n\\n"
        "Return the polished summary only."
    )
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    content = response.choices[0].message.content if response.choices else ""
    return (content or "").strip() or summary


def build_summary_prompts(
    *,
    form_data: dict,
    transcript: str,
) -> tuple[str, str]:
    """Build system/user prompts for five-paragraph narrative generation."""
    system_prompt = (
        "You are a professional arborist report writer. "
        "Write a clear, client-ready summary based on the provided extracted data "
        "and transcript. Use complete sentences, no bullet lists."
    )
    user_prompt = (
        "Write five short paragraphs in this exact order:\n"
        "1) Combine site factors with the tree details and assessment context.\n"
        "2) Target assessment with tree health & species profile. "
        "Explicitly name the targets and describe how far they are (e.g., within dripline, within 1x height).\n"
        "3) Crown & branches and trunk defects together.\n"
        "4) Roots & root collar with load factors.\n"
        "5) Overall risk, mitigation option(s), and expected residual risk after mitigation.\n\n"
        "Do not repeat customer contact info or job address; those belong in the letter header.\n"
        "Do not mention the assessor in the body; that belongs in the signature block.\n"
        "Use extracted form data as the primary source of truth (this reflects user-reviewed corrections).\n"
        "Use transcript only as supporting context to explain observations and values.\n"
        "Keep language conservative and factual for a legal document; do not overstate certainty.\n"
        "Do not add conclusions or recommendations not supported by form data/transcript.\n\n"
        "Extracted form data (JSON):\n"
        f"{json.dumps(form_data, ensure_ascii=True)}\n\n"
        "Transcript:\n"
        f"{transcript}\n"
    )
    return system_prompt, user_prompt


def generate_summary(
    *,
    form_data: dict,
    transcript: str,
) -> str:
    """Generate narrative summary from form data and transcript using OpenAI."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set")
    model = os.environ.get("TRAQ_OPENAI_SUMMARY_MODEL", "gpt-4o-mini")
    client = OpenAI(api_key=api_key)

    system_prompt, user_prompt = build_summary_prompts(
        form_data=form_data, transcript=transcript
    )

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    content = response.choices[0].message.content if response.choices else ""
    return (content or "").strip()


def _strip_state_zip_from_lines(lines: Sequence[str]) -> list[str]:
    """Strip state/ZIP suffix patterns from address lines."""
    cleaned: list[str] = []
    for line in lines:
        if not line.strip():
            continue
        cleaned.append(_strip_state_zip_from_line(line))
    return _compact_lines(cleaned)


def _strip_state_zip_from_line(line: str) -> str:
    """Strip state/ZIP tokens from a single line when detected."""
    text = " ".join(line.split()).strip()
    if not text:
        return text
    if "," in text:
        parts = [p.strip() for p in text.split(",") if p.strip()]
        if parts and _looks_like_state_zip(parts[-1]):
            parts = parts[:-1]
        text = ", ".join(parts)
    tokens = [t for t in text.split() if t]
    if len(tokens) >= 2 and _is_state(tokens[-2]) and _is_zip(tokens[-1]):
        text = " ".join(tokens[:-2])
    return text.strip()


def _split_address_line(address: str) -> tuple[str, str]:
    """Split a freeform address into street and city/state segments."""
    if not address:
        return "", ""
    parts = [p.strip() for p in address.split(",") if p.strip()]
    if len(parts) == 1:
        return parts[0], ""
    street = parts[0]
    city_state = ", ".join(parts[1:])
    city_state = _strip_state_zip_from_line(city_state)
    return street, city_state


def _build_user_address_block(profile: dict) -> list[str]:
    """Build sender address block from user profile fields."""
    street = _as_text(profile.get("correspondence_street"))
    city = _as_text(profile.get("correspondence_city"))
    state = _as_text(profile.get("correspondence_state"))
    phone = _as_text(profile.get("phone"))
    lines = []
    if street:
        lines.append(street)
    city_state = ", ".join([part for part in [city, state] if part])
    if city_state:
        lines.append(city_state)
    if phone:
        lines.append(phone)
    return _compact_lines(lines)


def _build_customer_address_block(job: dict) -> list[str]:
    """Build customer/recipient address block from job fields."""
    name = _as_text(job.get("billing_name"))
    contact = _as_text(job.get("billing_contact_name"))
    address = _as_text(job.get("billing_address") or job.get("job_address") or job.get("address"))
    phone = _as_text(job.get("job_phone"))
    street, city_state = _split_address_line(address)
    lines = [name, contact, street, city_state, phone]
    return _compact_lines(lines)


def _build_subject_line(job: dict, form_data: dict) -> str:
    """Build compact subject line from species and site location."""
    species = ""
    address = ""
    try:
        client_tree = form_data.get("data", {}).get("client_tree_details") if isinstance(form_data, dict) else None
        if client_tree:
            species = _as_text(client_tree.get("tree_species"))
            address = _as_text(client_tree.get("address_tree_location"))
    except Exception:
        pass
    if not address:
        address = _as_text(job.get("job_address") or job.get("address"))
    street, _ = _split_address_line(address)
    parts = []
    if species:
        parts.append(f"Tree species {species}")
    if street:
        parts.append(street)
    return ", ".join(parts).strip()
    return cleaned


def _looks_like_state_zip(value: str) -> bool:
    """Return True when value resembles state/ZIP suffix."""
    text = value.strip()
    if not text:
        return False
    if _is_zip(text):
        return True
    parts = [p for p in text.replace(",", " ").split() if p]
    if len(parts) == 2 and _is_state(parts[0]) and _is_zip(parts[1]):
        return True
    return False


def _strip_state_zip_from_subject(subject: str) -> str:
    """Strip state/ZIP components from subject line."""
    subject = " ".join(subject.splitlines()).strip()
    if not subject:
        return subject
    if "," in subject:
        parts = [p.strip() for p in subject.split(",") if p.strip()]
        if parts and _looks_like_state_zip(parts[-1]):
            parts = parts[:-1]
        subject = ", ".join(parts)
    tokens = [t for t in subject.split() if t]
    if len(tokens) >= 2 and _is_state(tokens[-2]) and _is_zip(tokens[-1]):
        subject = " ".join(tokens[:-2])
    return _truncate_one_line(subject)


def _truncate_one_line(value: str, limit: int = 80) -> str:
    """Truncate a string to one line with ellipsis."""
    if len(value) <= limit:
        return value
    return value[: limit - 1].rstrip() + "…"


def _is_state(value: str) -> bool:
    """Return True for 2-letter uppercase US-state-like token."""
    if len(value) != 2:
        return False
    return value.isalpha() and value.isupper()


def _is_zip(value: str) -> bool:
    """Return True for ZIP token lengths (5 or 9 digits)."""
    return value.isdigit() and (len(value) == 5 or len(value) == 9)


def _format_date(value: datetime | None = None) -> str:
    """Format date as 'Month DD, YYYY' for letter header."""
    dt = value or datetime.now()
    return dt.strftime("%B %d, %Y")
